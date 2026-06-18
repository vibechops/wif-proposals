#!/usr/bin/env python3
"""Generate an editable Word document from proposal content."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Google Flow x WIF India - Proposal.docx"

PLUM = RGBColor(0x3D, 0x1F, 0x3A)
EMBER = RGBColor(0xC6, 0x54, 0x3C)
INK = RGBColor(0x1A, 0x14, 0x18)


def set_run_font(run, size=11, bold=False, italic=False, color=INK, name="Calibri"):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_page_break(doc):
    doc.add_page_break()


def add_kicker(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    set_run_font(run, size=9, bold=True, color=EMBER)
    run.font.all_caps = True
    p.paragraph_format.space_after = Pt(4)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=22, bold=False, color=PLUM, name="Georgia")
        p.paragraph_format.space_after = Pt(14)
    else:
        set_run_font(run, size=14, bold=True, color=PLUM, name="Georgia")
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    return p


def add_body(doc, text, space_after=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=11)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.25
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25


def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=10, italic=True, color=RGBColor(0x5A, 0x52, 0x58))
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.line_spacing = 1.25


def add_comparison_table(doc, rows):
    table = doc.add_table(rows=1 + len(rows), cols=3)
    table.style = "Table Grid"
    headers = ["Program feature", "Option A", "Option B"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True, color=PLUM)

    for r_idx, (label, a, b) in enumerate(rows, start=1):
        mark = lambda on: "Yes" if on else "-"
        values = [label, mark(a), mark(b)]
        for c_idx, val in enumerate(values):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_run_font(run, size=10, bold=(c_idx == 0))

    doc.add_paragraph()


def add_roles_block(doc, tag, items):
    add_heading(doc, tag, level=2)
    for item in items:
        p = doc.add_paragraph()
        who = p.add_run(f"{item['who']}: ")
        set_run_font(who, size=10, bold=True, color=EMBER)
        what = p.add_run(item["what"])
        set_run_font(what, size=11)
        p.paragraph_format.space_after = Pt(6)


def add_timeline(doc, option_tag, intro, phases):
    add_kicker(doc, "Program schedule")
    add_heading(doc, "Timeline")
    p = doc.add_paragraph()
    tag = p.add_run(f"{option_tag} · ")
    set_run_font(tag, size=11, bold=True, color=EMBER)
    intro_run = p.add_run(intro)
    set_run_font(intro_run, size=11)
    p.paragraph_format.space_after = Pt(12)

    for phase in phases:
        add_heading(doc, f"{phase['label']} ({phase['range']})", level=2)
        if "cities" in phase:
            add_body(doc, "Cities: " + ", ".join(phase["cities"]), space_after=4)
            if phase.get("note"):
                add_note(doc, phase["note"])
        if "items" in phase:
            add_bullets(doc, phase["items"])


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Page 1 - Cover
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Google Flow × WIF India")
    set_run_font(run, size=28, color=PLUM, name="Georgia")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Partnership Proposal")
    set_run_font(run, size=14, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("A 7-city roadshow across India's top film institutes")
    set_run_font(run, size=12, italic=True, color=RGBColor(0x5A, 0x52, 0x58))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2026")
    set_run_font(run, size=11, color=RGBColor(0x5A, 0x52, 0x58))

    add_page_break(doc)

    # Page 2 - In one sentence
    add_kicker(doc, "The proposal, in one sentence")
    add_heading(doc, "In one sentence")
    add_body(
        doc,
        "WIF India proposes a Google + WIF roadshow across 7 cities, led by a Google Flow expert.",
        space_after=14,
    )
    add_bullets(
        doc,
        [
            "The partnership: Google and WIF India, together on the road.",
            "The reach: Seven cities, India's top film institutes.",
            "The expertise: a Google Flow expert leads every workshop and demo.",
        ],
    )
    add_note(doc, "Google chooses one of two program levels: events only, or events plus Filmathon and awards.")

    add_page_break(doc)

    # Page 3 - Why this program
    add_kicker(doc, "Why now")
    add_heading(doc, "Why this program")
    add_bullets(
        doc,
        [
            "India's next generation of filmmakers is entering an industry where AI tools are becoming part of the creative stack, but hands-on access at film schools and creator communities is still limited.",
            "Google is already investing in AI filmmaking education globally (Flow Sessions, Sundance Institute partnership). India needs the same activation layer: real workshops, real creators, real films, not just product demos.",
            "A 7-city roadshow puts Google's Flow expert inside India's top film institutes and connects them to working professionals, exactly where new creative habits form.",
        ],
    )

    add_page_break(doc)

    # Page 4 - Why WIF India
    add_kicker(doc, "Partner credentials")
    add_heading(doc, "Why WIF India")

    sections = [
        (
            "Leadership",
            [
                "Founded by Guneet Monga Kapoor, Oscar-winning producer (The Elephant Whisperers)",
                "Sister organization to WIF LA, a globally recognized advocacy body for women in film",
                "Rabia Chopra, Director: leads program management, partnerships, and industry engagement on the ground",
            ],
        ),
        (
            "Community and reach",
            [
                "3,500+ member national community across film, TV, and digital media",
                "Free membership; spans students, emerging talent, and working professionals",
                "The Collective Database: curated network of women and gender-diverse film professionals across India",
            ],
        ),
        (
            "Why WIF fits this roadshow",
            [
                "Existing relationships with top film institutes across India",
                "Track record hosting networking events, workshops, and masterclasses nationally",
                "Credible bridge between Google's product story and India's film school + industry audience",
            ],
        ),
    ]
    for title, items in sections:
        add_heading(doc, title, level=2)
        add_bullets(doc, items)

    add_page_break(doc)

    # Page 5 - Two options
    add_kicker(doc, "Choose a track")
    add_heading(doc, "Pick your program")
    comparison_rows = [
        ("7-city roadshow at India's top film colleges", True, True),
        ("Google Flow expert-led workshop", True, True),
        ("Students + professionals networking", True, True),
        ("16-week program", True, True),
        ("Online workshops", True, True),
        ("National short film call for submissions", False, True),
        ("WIF script curation + filmmaker selection", False, True),
        ("AI Filmathon (14-day production in Flow)", False, True),
        ("Mumbai awards ceremony (Top 3)", False, True),
        ("1-year free Google Flow subscription per winner", False, True),
    ]
    add_comparison_table(doc, comparison_rows)
    add_body(doc, "Option A: Events. College events across 7 cities. Stops after the roadshow.")
    add_body(doc, "Option B: Events + Filmathon. Everything in A, plus call for submissions, Filmathon, and Mumbai awards.")

    add_page_break(doc)

    # Page 6 - Who does what
    add_kicker(doc, "Division of work")
    add_heading(doc, "Who does what")
    add_roles_block(
        doc,
        "Option A",
        [
            {"who": "Google", "what": "Flow expert on the 7-city roadshow"},
            {"who": "WIF India", "what": "Venues, event production, and audience"},
        ],
    )
    add_roles_block(
        doc,
        "Option B",
        [
            {"who": "Google", "what": "Flow expert on the 7-city roadshow"},
            {"who": "WIF India", "what": "Venues, event production, and audience"},
            {"who": "WIF India", "what": "National short film call for submissions: curation and filmmaker selection"},
            {"who": "WIF India", "what": "Filmathon operations and Flow credits"},
            {"who": "Google", "what": "Flow expert at Mumbai awards"},
            {"who": "WIF India", "what": "Mumbai awards event"},
        ],
    )

    add_page_break(doc)

    # Page 7 - Timeline A
    add_timeline(
        doc,
        "Option A",
        "16 weeks · Roadshow across 7 cities, then wrap.",
        [
            {
                "label": "Prep",
                "range": "Weeks 1–4",
                "items": [
                    "Confirm venues, align with Google expert",
                    "Open registrations, lock production",
                ],
            },
            {
                "label": "Roadshow",
                "range": "Weeks 5–12",
                "cities": ["Mumbai", "Pune", "Kolkata", "Delhi", "Chennai", "Hyderabad", "Bangalore"],
            },
            {
                "label": "Wrap",
                "range": "Weeks 13–16",
                "items": [
                    "Synthesize roadshow learnings",
                    "Report and content package to Google",
                ],
            },
        ],
    )

    add_page_break(doc)

    # Page 8 - Timeline B
    add_timeline(
        doc,
        "Option B",
        "16 weeks · Roadshow + call for submissions, Filmathon, and Mumbai awards.",
        [
            {
                "label": "Prep",
                "range": "Weeks 1–4",
                "items": [
                    "Confirm venues, align with Google expert",
                    "Registrations + call for submissions brief",
                ],
            },
            {
                "label": "Roadshow + call for submissions",
                "range": "Weeks 5–12",
                "cities": ["Mumbai", "Pune", "Kolkata", "Delhi", "Chennai", "Hyderabad", "Bangalore"],
                "note": "National call for submissions opens Week 5, closes Week 12",
            },
            {
                "label": "Filmathon + awards",
                "range": "Weeks 13–16",
                "items": [
                    "WIF and Google select filmmakers",
                    "14-day AI Filmathon in Flow",
                    "Mumbai awards · Top 3",
                ],
            },
        ],
    )

    add_page_break(doc)

    # Page 9 - Budget
    add_kicker(doc, "The investment")
    add_heading(doc, "Budget")
    p = doc.add_paragraph()
    label = p.add_run("WIF India program fee")
    set_run_font(label, size=10, bold=True, color=EMBER)
    p.paragraph_format.space_after = Pt(8)

    for tag, amt, desc in [
        ("Option A", "USD $100,000", "Roadshow across 7 cities, then wrap."),
        ("Option B", "USD $150,000", "Roadshow + national call for submissions, Filmathon, and Mumbai awards."),
    ]:
        tp = doc.add_paragraph()
        tag_run = tp.add_run(f"{tag}   ")
        set_run_font(tag_run, size=10, bold=True, color=PLUM)
        amt_run = tp.add_run(amt)
        set_run_font(amt_run, size=22, color=PLUM, name="Georgia")
        tp.paragraph_format.space_after = Pt(2)
        plus_p = doc.add_paragraph()
        plus_run = plus_p.add_run("+ production & other expenses")
        set_run_font(plus_run, size=10, bold=True, color=EMBER)
        plus_p.paragraph_format.space_after = Pt(2)
        dp = doc.add_paragraph()
        d_run = dp.add_run(desc)
        set_run_font(d_run, size=10, color=RGBColor(0x5A, 0x52, 0x58))
        dp.paragraph_format.space_after = Pt(10)

    add_body(
        doc,
        "A WIF India program fee per option, plus production and other expenses scoped to the option you choose. WIF plans, produces, and manages the 16-week partnership end to end.",
    )
    add_bullets(
        doc,
        [
            "Event production across 7 cities: venues, logistics, and audience",
            "On-ground program management, partnerships, and industry engagement",
            "National short film call for submissions: curation and filmmaker coordination (Option B)",
            "Filmathon operations and participant Flow access (Option B)",
            "Mumbai awards ceremony (Option B)",
        ],
    )
    add_note(
        doc,
        "Google provides the Flow expert for workshops and the awards event. Top 3 winners each receive a 1-year free Google Flow subscription.",
    )

    add_page_break(doc)

    # Page 10 - The ask
    add_kicker(doc, "The decision")
    add_heading(doc, "The ask")
    add_body(
        doc,
        "At its core, this is a Google + WIF roadshow across 7 cities in India over 16 weeks.",
    )
    add_heading(doc, "What we ask of Google", level=2)
    add_bullets(
        doc,
        [
            "Pick Option A or Option B",
            "Provide Google's Flow expert for the 7-city roadshow",
            "Top 3 winners each receive a 1-year free Google Flow subscription",
        ],
    )

    add_page_break(doc)

    # Page 11 - The vision (closing)
    add_kicker(doc, "The opportunity")
    p = doc.add_paragraph()
    run = p.add_run("Google AI × WIF India")
    set_run_font(run, size=13, bold=True, color=EMBER)
    p.paragraph_format.space_after = Pt(18)

    p = doc.add_paragraph()
    run = p.add_run(
        "Empowering the storytellers who will define the next decade of global cinema."
    )
    set_run_font(run, size=26, color=PLUM, name="Georgia")
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.space_after = Pt(24)

    p = doc.add_paragraph()
    run = p.add_run("7 cities  ·  16 weeks  ·  3,500+ creators in reach")
    set_run_font(run, size=11, bold=True, color=INK)
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("Let's put Flow in their hands first.")
    set_run_font(run, size=14, italic=True, color=EMBER, name="Georgia")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
